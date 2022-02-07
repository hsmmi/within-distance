import os
import numpy as np
import pandas as pd
from normalization import clipping, log_scaling, range_min_to_max,\
     zero_mean_unit_variance


def read_dataset(file, atr):
    if (type(atr) == int):
        with open(os.path.join(os.path.dirname(__file__), file), 'r') as f:
            return list(map(
                lambda x: float(x.split(', ')[atr]), f.read().splitlines()))
    else:
        with open(os.path.join(os.path.dirname(__file__), file), 'r') as f:
            return list(map(
                lambda x: [float(i) for i in (x.split(', ')[atr[0]:atr[1]])],
                f.read().splitlines()))


def read_dataset_with_pandas(file, atr=None):
    colName = pd.read_csv(
        os.path.join(os.path.dirname(__file__), file), nrows=0).columns
    if (type(atr) == int):
        colName = [colName[atr]]
    elif(atr is not None):
        colName = colName[atr[0]:atr[1]]
    data = pd.read_csv(
        os.path.join(os.path.dirname(__file__), file), usecols=colName)

    return colName, data


def dataframe_to_docx_table(header, data, file, doc=None, save=1):
    """
    Read header and data
    If you gave if doc it add header and data to it and return it
    If you gave it save=0 it will not be save doc
    Return doc include header and data
    """
    import docx
    if(doc is None):
        doc = docx.Document()
    doc.add_heading(header, 1)

    table = doc.add_table(rows=len(data.index)+1, cols=len(data.columns)+1)

    for j in range(len(data.columns)):
        table.cell(0, j+1).text = f'{data.columns[j]}'

    for i in range(len(data.index)):
        table.cell(i+1, 0).text = f'{data.index[i]}'
        for j in range(len(data.columns)):
            table.cell(i+1, j+1).text = f'{data.iat[i, j]}'
    table.style = 'Table Grid'
    if(save):
        doc.save(file)
    return doc


def string_to_dataframe(string):
    from io import StringIO
    data = StringIO(string)
    return pd.read_csv(data)


def read_dataset_to_X_and_y(
        file, range_feature=None, range_label=None, normalization=None,
        min_value=None, max_value=None, add_x0=False, shuffle=False,
        about_nan=None):
    """
    Read the attribute(range_atr) that you want and put X0 = 1 and thoes
    attribute of all samples in X and all samples label in y
    normalization:
    .   by default is None and can be "z_score", "scaling", "clipping"
        or "log_scaling"
    .   for "scaling", "clipping" must set min_value and max_value
    Return X and y as nparray
    """
    import numpy as np
    col_name, data = read_dataset_with_pandas(file)
    if(about_nan == 'delete'):
        data.dropna(inplace=True)

    number_of_attribute = len(col_name)
    for col in col_name:
        if(data[col].dtype != int and data[col].dtype != float):
            data[col] = pd.factorize(data[col])[0]

    data = data.to_numpy()

    if(shuffle is True):
        np.random.shuffle(data)

    if(range_feature is None):
        range_feature = (0, number_of_attribute-1)
    if(range_label is None):
        range_label = (number_of_attribute-1, number_of_attribute)

    feature = np.array(list(map(
        lambda x: x[range_feature[0]:range_feature[1]], data)))
    label = np.array(list(map(
        lambda x: x[range_label[0]:range_label[1]], data)))

    if(about_nan == 'class_mean'):
        feature = feature.astype(float)
        diffrent_label = np.unique(label)
        number_of_feature = feature.shape[1]
        number_of_sample = feature.shape[0]
        for a_label in diffrent_label:
            class_label = feature[(label == a_label).flatten()]
            for a_feature in range(number_of_feature):
                mean_feature_label = np.nanmean(class_label[:, a_feature])
                for a_sample in range(number_of_sample):
                    if np.isnan(feature[a_sample, a_feature]):
                        feature[a_sample, a_feature] = mean_feature_label

    if(normalization is not None):
        if(normalization == 'z_score'):
            feature = zero_mean_unit_variance(feature)
        elif(normalization == 'scaling'):
            feature = range_min_to_max(feature, min_value, max_value)
        elif(normalization == 'clipping'):
            feature = clipping(feature, min_value, max_value)
        elif(normalization == 'logScaling'):
            feature = log_scaling(feature)
        else:
            print(
                'method should be "z_score", "scaling", "clipping" or'
                '"logScaling"')
            return

    return feature, label


def nparray_to_csv(file: str, input: np.ndarray, decimal: int) -> None:
    path = os.path.join(os.path.dirname(__file__), file)
    pd.DataFrame(np.round(input, decimal)).to_csv(path)


def print_array_with_dataframe(array):
    print(pd.DataFrame(array))
